"""文档导出服务：查询结果与数据库文档，支持 csv/json/excel/word/markdown/html。"""
import csv
import io
import json
import threading
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from fastapi import HTTPException

from ..adapters import AdapterError
from ..config import EXPORTS_DIR
from ..models import DataSource
from .sql_service import build_adapter, classify_statement, split_statements

_pool = ThreadPoolExecutor(max_workers=2)
_tasks: dict = {}
_lock = threading.Lock()


def _serialize(value):
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return value


def _run_query(ds: DataSource, sql: str) -> tuple[list, list]:
    statements = split_statements(sql)
    if not statements:
        raise HTTPException(status_code=400, detail="SQL 为空")
    if any(classify_statement(s) != "READ" for s in statements):
        raise HTTPException(status_code=400, detail="导出仅支持只读查询")
    adapter = build_adapter(ds)
    result = None
    for stmt in statements:
        res = adapter.execute(stmt)
        if res.get("is_query"):
            result = res
    if result is None:
        raise HTTPException(status_code=400, detail="查询无结果集")
    return result["columns"], result["rows"]


# ---------- 结果导出 ----------
def _to_csv(columns: list, rows: list) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(columns)
    for row in rows:
        writer.writerow([_serialize(v) for v in row])
    return ("\ufeff" + buf.getvalue()).encode("utf-8")


def _to_json(columns: list, rows: list) -> bytes:
    data = [dict(zip(columns, [_serialize(v) for v in row])) for row in rows]
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def _to_excel(columns: list, rows: list, sheet_name: str = "结果") -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "结果"
    ws.append([str(c) for c in columns])
    for row in rows:
        ws.append([_serialize(v) for v in row])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_result_file(ds: DataSource, sql: str, fmt: str, sheet_name: str = "结果") -> tuple[bytes, str]:
    columns, rows = _run_query(ds, sql)
    fmt = (fmt or "csv").lower()
    if fmt == "csv":
        return _to_csv(columns, rows), "result.csv"
    if fmt == "json":
        return _to_json(columns, rows), "result.json"
    if fmt == "excel" or fmt == "xlsx":
        return _to_excel(columns, rows, sheet_name), "result.xlsx"
    raise HTTPException(status_code=400, detail=f"不支持的导出格式: {fmt}")


# ---------- 数据库文档导出 ----------
def _collect_metadata(ds: DataSource, tables: list[str] | None) -> list[dict]:
    adapter = build_adapter(ds)
    names = tables or adapter.get_tables()
    docs = []
    for table in names[:200]:
        try:
            columns = adapter.get_table_columns(table)
            indexes = adapter.get_table_indexes(table)
            ddl = adapter.get_table_ddl(table) if hasattr(adapter, "get_table_ddl") else ""
            docs.append({"schema": "", "table": table, "columns": columns, "indexes": indexes, "ddl": ddl})
        except AdapterError:
            continue
    return docs


def _to_markdown(docs: list[dict]) -> str:
    lines = ["# 数据库结构文档", ""]
    for doc in docs:
        lines.append(f"## 表：{doc['table']}", "")
        lines.append("| 列名 | 类型 | 可空 | 默认值 | 主键 | 自增 | 注释 |")
        lines.append("| --- | --- | --- | --- | --- | --- | --- |")
        for c in doc["columns"]:
            lines.append(
                f"| {c['name']} | {c['data_type']} | {'是' if c['nullable'] else '否'} | {c.get('default') or ''} | {'✓' if c['primary_key'] else ''} | {'✓' if c.get('auto_increment') else ''} | {c.get('comment') or ''} |"
            )
        if doc["ddl"]:
            lines.append("", "```sql", doc["ddl"], "```")
        lines.append("")
    return "\n".join(lines)


def _to_html(docs: list[dict]) -> str:
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'><title>数据库结构文档</title>",
        "<style>body{font-family:system-ui;margin:24px}table{border-collapse:collapse;margin:12px 0;width:100%}th,td{border:1px solid #ddd;padding:6px 10px;font-size:13px}th{background:#f5f5f5}pre{background:#f6f8fa;padding:12px;border-radius:6px;overflow:auto}</style></head><body>",
        "<h1>数据库结构文档</h1>",
    ]
    for doc in docs:
        parts.append(f"<h2>表：{doc['table']}</h2>")
        parts.append("<table><tr><th>列名</th><th>类型</th><th>可空</th><th>默认值</th><th>主键</th><th>自增</th><th>注释</th></tr>")
        for c in doc["columns"]:
            parts.append(
                f"<tr><td>{c['name']}</td><td>{c['data_type']}</td><td>{'是' if c['nullable'] else '否'}</td><td>{c.get('default') or ''}</td><td>{'✓' if c['primary_key'] else ''}</td><td>{'✓' if c.get('auto_increment') else ''}</td><td>{c.get('comment') or ''}</td></tr>"
            )
        parts.append("</table>")
        if doc["ddl"]:
            parts.append(f"<pre>{doc['ddl']}</pre>")
    parts.append("</body></html>")
    return "\n".join(parts)


def _to_word(docs: list[dict]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("数据库结构文档", 0)
    for doc in docs:
        document.add_heading(f"表：{doc['table']}", level=1)
        table = document.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(["列名", "类型", "可空", "默认值", "主键", "自增", "注释"]):
            hdr[i].text = name
        for c in doc["columns"]:
            row = table.add_row().cells
            row[0].text = str(c["name"])
            row[1].text = str(c["data_type"])
            row[2].text = "是" if c["nullable"] else "否"
            row[3].text = str(c.get("default") or "")
            row[4].text = "✓" if c["primary_key"] else ""
            row[5].text = "✓" if c.get("auto_increment") else ""
            row[6].text = str(c.get("comment") or "")
        if doc["ddl"]:
            document.add_paragraph("DDL：")
            document.add_paragraph(doc["ddl"], style="Normal")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _to_excel_doc(docs: list[dict]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "总览"
    ws.append(["表名", "列数", "DDL"])
    for doc in docs:
        ws.append([doc["table"], len(doc["columns"]), doc["ddl"]])
    for doc in docs:
        sheet = wb.create_sheet(title=(doc["table"])[:31])
        sheet.append(["列名", "类型", "可空", "默认值", "主键", "自增", "注释"])
        for c in doc["columns"]:
            sheet.append([c["name"], c["data_type"], "是" if c["nullable"] else "否", c.get("default") or "", "✓" if c["primary_key"] else "", "✓" if c.get("auto_increment") else "", c.get("comment") or ""])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_database_file(ds: DataSource, tables: list[str] | None, fmt: str, include_ddl: bool = True) -> tuple[bytes, str]:
    docs = _collect_metadata(ds, tables)
    fmt = (fmt or "word").lower()
    if fmt == "markdown" or fmt == "md":
        return _to_markdown(docs).encode("utf-8"), "database_doc.md"
    if fmt == "html":
        return _to_html(docs).encode("utf-8"), "database_doc.html"
    if fmt == "word" or fmt == "docx":
        return _to_word(docs), "database_doc.docx"
    if fmt == "excel" or fmt == "xlsx":
        return _to_excel_doc(docs), "database_doc.xlsx"
    raise HTTPException(status_code=400, detail=f"不支持的导出格式: {fmt}")


# ---------- 异步任务 ----------
def start_database_export(ds: DataSource, tables: list[str] | None, fmt: str, include_ddl: bool = True) -> str:
    task_id = uuid.uuid4().hex
    with _lock:
        _tasks[task_id] = {"status": "running", "progress": 0, "file_name": None, "path": None, "error": None, "created_at": time.time()}

    def run():
        try:
            data, name = export_database_file(ds, tables, fmt, include_ddl)
            path = EXPORTS_DIR / f"{task_id}_{name}"
            path.write_bytes(data)
            with _lock:
                _tasks[task_id].update({"status": "done", "progress": 100, "file_name": name, "path": str(path)})
        except Exception as exc:  # noqa: BLE001
            with _lock:
                _tasks[task_id].update({"status": "failed", "error": str(exc)})

    _pool.submit(run)
    return task_id


def get_task_status(task_id: str) -> dict:
    with _lock:
        task = _tasks.get(task_id)
        if task is None:
            raise HTTPException(status_code=404, detail="导出任务不存在")
        result = dict(task)
        if result["status"] == "done":
            result["download_url"] = f"/api/export/database/download/{task_id}"
        return result


def get_task_file(task_id: str) -> Path:
    with _lock:
        task = _tasks.get(task_id)
        if task is None or task["status"] != "done" or not task.get("path"):
            raise HTTPException(status_code=404, detail="导出文件不存在或未完成")
        path = Path(task["path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="导出文件不存在")
    return path
"""文档导出服务：查询结果与数据库文档，支持 csv/json/excel/word/markdown/html。"""
import csv
import io
import json
import threading
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from fastapi import HTTPException

from ..adapters import AdapterError
from ..config import EXPORTS_DIR
from ..models import DataSource
from .sql_service import build_adapter, classify_statement, split_statements

_pool = ThreadPoolExecutor(max_workers=2)
_tasks: dict = {}
_lock = threading.Lock()


def _serialize(value):
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return value


def _run_query(ds: DataSource, sql: str) -> tuple[list, list]:
    statements = split_statements(sql)
    if not statements:
        raise HTTPException(status_code=400, detail="SQL 为空")
    if any(classify_statement(s) != "READ" for s in statements):
        raise HTTPException(status_code=400, detail="导出仅支持只读查询")
    adapter = build_adapter(ds)
    result = None
    for stmt in statements:
        res = adapter.execute(stmt)
        if res.get("is_query"):
            result = res
    if result is None:
        raise HTTPException(status_code=400, detail="查询无结果集")
    return result["columns"], result["rows"]


# ---------- 结果导出 ----------
def _to_csv(columns: list, rows: list) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(columns)
    for row in rows:
        writer.writerow([_serialize(v) for v in row])
    return ("\ufeff" + buf.getvalue()).encode("utf-8")


def _to_json(columns: list, rows: list) -> bytes:
    data = [dict(zip(columns, [_serialize(v) for v in row])) for row in rows]
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def _to_excel(columns: list, rows: list, sheet_name: str = "结果") -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "结果"
    ws.append([str(c) for c in columns])
    for row in rows:
        ws.append([_serialize(v) for v in row])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_result_file(ds: DataSource, sql: str, fmt: str, sheet_name: str = "结果") -> tuple[bytes, str]:
    columns, rows = _run_query(ds, sql)
    fmt = (fmt or "csv").lower()
    if fmt == "csv":
        return _to_csv(columns, rows), "result.csv"
    if fmt == "json":
        return _to_json(columns, rows), "result.json"
    if fmt == "excel" or fmt == "xlsx":
        return _to_excel(columns, rows, sheet_name), "result.xlsx"
    raise HTTPException(status_code=400, detail=f"不支持的导出格式: {fmt}")


# ---------- 数据库文档导出 ----------
def _collect_metadata(ds: DataSource, tables: list[str] | None) -> list[dict]:
    adapter = build_adapter(ds)
    names = tables or adapter.get_tables()
    docs = []
    for table in names[:200]:
        try:
            columns = adapter.get_table_columns(table)
            indexes = adapter.get_table_indexes(table)
            ddl = adapter.get_table_ddl(table) if hasattr(adapter, "get_table_ddl") else ""
            docs.append({"schema": "", "table": table, "columns": columns, "indexes": indexes, "ddl": ddl})
        except AdapterError:
            continue
    return docs


def _to_markdown(docs: list[dict]) -> str:
    lines = ["# 数据库结构文档", ""]
    for doc in docs:
        lines.append(f"## 表：{doc['table']}")
        lines.append("")
        lines.append("| 列名 | 类型 | 可空 | 默认值 | 主键 | 自增 | 注释 |")
        lines.append("| --- | --- | --- | --- | --- | --- | --- |")
        for c in doc["columns"]:
            lines.append(
                f"| {c['name']} | {c['data_type']} | {'是' if c['nullable'] else '否'} | {c.get('default') or ''} | {'✓' if c['primary_key'] else ''} | {'✓' if c.get('auto_increment') else ''} | {c.get('comment') or ''} |"
            )
        if doc["ddl"]:
            lines.append("")
            lines.append("```sql")
            lines.append(doc["ddl"])
            lines.append("```")
        lines.append("")
    return "\n".join(lines)


def _to_html(docs: list[dict]) -> str:
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'><title>数据库结构文档</title>",
        "<style>body{font-family:system-ui;margin:24px}table{border-collapse:collapse;margin:12px 0;width:100%}th,td{border:1px solid #ddd;padding:6px 10px;font-size:13px}th{background:#f5f5f5}pre{background:#f6f8fa;padding:12px;border-radius:6px;overflow:auto}</style></head><body>",
        "<h1>数据库结构文档</h1>",
    ]
    for doc in docs:
        parts.append(f"<h2>表：{doc['table']}</h2>")
        parts.append("<table><tr><th>列名</th><th>类型</th><th>可空</th><th>默认值</th><th>主键</th><th>自增</th><th>注释</th></tr>")
        for c in doc["columns"]:
            parts.append(
                f"<tr><td>{c['name']}</td><td>{c['data_type']}</td><td>{'是' if c['nullable'] else '否'}</td><td>{c.get('default') or ''}</td><td>{'✓' if c['primary_key'] else ''}</td><td>{'✓' if c.get('auto_increment') else ''}</td><td>{c.get('comment') or ''}</td></tr>"
            )
        parts.append("</table>")
        if doc["ddl"]:
            parts.append(f"<pre>{doc['ddl']}</pre>")
    parts.append("</body></html>")
    return "\n".join(parts)


def _to_word(docs: list[dict]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("数据库结构文档", 0)
    for doc in docs:
        document.add_heading(f"表：{doc['table']}", level=1)
        table = document.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(["列名", "类型", "可空", "默认值", "主键", "自增", "注释"]):
            hdr[i].text = name
        for c in doc["columns"]:
            row = table.add_row().cells
            row[0].text = str(c["name"])
            row[1].text = str(c["data_type"])
            row[2].text = "是" if c["nullable"] else "否"
            row[3].text = str(c.get("default") or "")
            row[4].text = "✓" if c["primary_key"] else ""
            row[5].text = "✓" if c.get("auto_increment") else ""
            row[6].text = str(c.get("comment") or "")
        if doc["ddl"]:
            document.add_paragraph("DDL：")
            document.add_paragraph(doc["ddl"], style="Normal")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _to_excel_doc(docs: list[dict]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "总览"
    ws.append(["表名", "列数", "DDL"])
    for doc in docs:
        ws.append([doc["table"], len(doc["columns"]), doc["ddl"]])
    for doc in docs:
        sheet = wb.create_sheet(title=(doc["table"])[:31])
        sheet.append(["列名", "类型", "可空", "默认值", "主键", "自增", "注释"])
        for c in doc["columns"]:
            sheet.append([c["name"], c["data_type"], "是" if c["nullable"] else "否", c.get("default") or "", "✓" if c["primary_key"] else "", "✓" if c.get("auto_increment") else "", c.get("comment") or ""])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_database_file(ds: DataSource, tables: list[str] | None, fmt: str, include_ddl: bool = True) -> tuple[bytes, str]:
    docs = _collect_metadata(ds, tables)
    fmt = (fmt or "word").lower()
    if fmt == "markdown" or fmt == "md":
        return _to_markdown(docs).encode("utf-8"), "database_doc.md"
    if fmt == "html":
        return _to_html(docs).encode("utf-8"), "database_doc.html"
    if fmt == "word" or fmt == "docx":
        return _to_word(docs), "database_doc.docx"
    if fmt == "excel" or fmt == "xlsx":
        return _to_excel_doc(docs), "database_doc.xlsx"
    raise HTTPException(status_code=400, detail=f"不支持的导出格式: {fmt}")


# ---------- 异步任务 ----------
def start_database_export(ds: DataSource, tables: list[str] | None, fmt: str, include_ddl: bool = True) -> str:
    task_id = uuid.uuid4().hex
    with _lock:
        _tasks[task_id] = {"status": "running", "progress": 0, "file_name": None, "path": None, "error": None, "created_at": time.time()}

    def run():
        try:
            data, name = export_database_file(ds, tables, fmt, include_ddl)
            path = EXPORTS_DIR / f"{task_id}_{name}"
            path.write_bytes(data)
            with _lock:
                _tasks[task_id].update({"status": "done", "progress": 100, "file_name": name, "path": str(path)})
        except Exception as exc:  # noqa: BLE001
            with _lock:
                _tasks[task_id].update({"status": "failed", "error": str(exc)})

    _pool.submit(run)
    return task_id


def get_task_status(task_id: str) -> dict:
    with _lock:
        task = _tasks.get(task_id)
        if task is None:
            raise HTTPException(status_code=404, detail="导出任务不存在")
        result = dict(task)
        if result["status"] == "done":
            result["download_url"] = f"/api/export/database/download/{task_id}"
        return result


def get_task_file(task_id: str) -> Path:
    with _lock:
        task = _tasks.get(task_id)
        if task is None or task["status"] != "done" or not task.get("path"):
            raise HTTPException(status_code=404, detail="导出文件不存在或未完成")
        path = Path(task["path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="导出文件不存在")
    return path
